"""
IntelliSafe - File Parser Module
Handles extraction of text from various file formats
"""

import logging
from pathlib import Path
from typing import Dict, Tuple, Optional

logger = logging.getLogger(__name__)


class PDFParser:
    """Parse PDF files using PyMuPDF"""
    
    @staticmethod
    def parse(file_path: Path) -> Dict:
        """
        Extract text and metadata from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with text content and metadata
        """
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text = ""
            metadata = doc.metadata
            page_count = len(doc)
            
            # Extract text from all pages
            for page_num, page in enumerate(doc):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.get_text()
            
            doc.close()
            
            logger.info(f"Successfully parsed PDF: {file_path.name} ({page_count} pages)")
            
            return {
                'status': 'success',
                'content': text,
                'metadata': {
                    'filename': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'pages': page_count,
                    'author': metadata.get('author', 'Unknown'),
                    'title': metadata.get('title', 'Unknown'),
                    'created': metadata.get('creationDate', 'Unknown'),
                },
                'format': 'pdf'
            }
        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return {
                'status': 'error',
                'error': 'PyMuPDF not installed',
                'message': 'Install PyMuPDF with: pip install PyMuPDF'
            }
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path.name}: {str(e)}")
            return {
                'status': 'error',
                'error': str(e),
                'file': file_path.name
            }


class DOCXParser:
    """Parse DOCX files using python-docx"""
    
    @staticmethod
    def parse(file_path: Path) -> Dict:
        """
        Extract text and metadata from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with text content and metadata
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
            
            doc = Document(file_path)
            text = ""
            paragraph_count = 0
            
            # Extract text from all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables if any
            if doc.tables:
                text += "\n--- Tables ---\n"
                for table_idx, table in enumerate(doc.tables):
                    text += f"\nTable {table_idx + 1}:\n"
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        text += " | ".join(row_text) + "\n"
            
            # Get core properties
            core_props = doc.core_properties
            
            logger.info(f"Successfully parsed DOCX: {file_path.name} ({paragraph_count} paragraphs)")
            
            return {
                'status': 'success',
                'content': text,
                'metadata': {
                    'filename': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'paragraphs': paragraph_count,
                    'tables': len(doc.tables),
                    'author': core_props.author or 'Unknown',
                    'title': core_props.title or 'Unknown',
                    'subject': core_props.subject or 'Unknown',
                },
                'format': 'docx'
            }
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return {
                'status': 'error',
                'error': 'python-docx not installed',
                'message': 'Install python-docx with: pip install python-docx'
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path.name}: {str(e)}")
            return {
                'status': 'error',
                'error': str(e),
                'file': file_path.name
            }


class TXTParser:
    """Parse plain text files"""
    
    @staticmethod
    def parse(file_path: Path) -> Dict:
        """
        Extract text from plain text file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Dictionary with text content and metadata
        """
        try:
            # Try to detect encoding
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    encoding_used = encoding
                    break
                except (UnicodeDecodeError, AttributeError):
                    continue
            
            if text is None:
                raise ValueError("Could not decode file with any supported encoding")
            
            line_count = len(text.split('\n'))
            word_count = len(text.split())
            
            logger.info(f"Successfully parsed TXT: {file_path.name} ({line_count} lines, encoding: {encoding_used})")
            
            return {
                'status': 'success',
                'content': text,
                'metadata': {
                    'filename': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'lines': line_count,
                    'words': word_count,
                    'encoding': encoding_used,
                },
                'format': 'txt'
            }
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path.name}: {str(e)}")
            return {
                'status': 'error',
                'error': str(e),
                'file': file_path.name
            }


class SQLParser:
    """Parse SQL files"""
    
    @staticmethod
    def parse(file_path: Path) -> Dict:
        """
        Extract SQL queries from SQL file
        
        Args:
            file_path: Path to SQL file
            
        Returns:
            Dictionary with SQL content and metadata
        """
        try:
            # Read SQL file
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Count queries (simple heuristic: count semicolons)
            queries = [q.strip() for q in content.split(';') if q.strip()]
            query_count = len(queries)
            
            # Count SQL keywords
            keywords = ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'DROP', 'CREATE', 'ALTER']
            keyword_count = {}
            for keyword in keywords:
                count = content.upper().count(keyword)
                if count > 0:
                    keyword_count[keyword] = count
            
            logger.info(f"Successfully parsed SQL: {file_path.name} ({query_count} queries)")
            
            return {
                'status': 'success',
                'content': content,
                'metadata': {
                    'filename': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'queries': query_count,
                    'keywords': keyword_count,
                },
                'format': 'sql'
            }
        except Exception as e:
            logger.error(f"Error parsing SQL {file_path.name}: {str(e)}")
            return {
                'status': 'error',
                'error': str(e),
                'file': file_path.name
            }


class FileParser:
    """Main file parser that routes to appropriate parser"""
    
    PARSERS = {
        '.pdf': PDFParser,
        '.docx': DOCXParser,
        '.txt': TXTParser,
        '.sql': SQLParser,
    }
    
    @staticmethod
    def parse(file_path: Path) -> Dict:
        """
        Parse file based on extension
        
        Args:
            file_path: Path to file to parse
            
        Returns:
            Dictionary with parsed content and metadata
        """
        extension = file_path.suffix.lower()
        
        if extension not in FileParser.PARSERS:
            logger.warning(f"Unsupported file format: {extension}")
            return {
                'status': 'error',
                'error': f'Unsupported file format: {extension}',
                'file': file_path.name
            }
        
        parser_class = FileParser.PARSERS[extension]
        logger.info(f"Parsing {extension} file: {file_path.name}")
        
        return parser_class.parse(file_path)
    
    @staticmethod
    def is_supported(file_path: Path) -> bool:
        """Check if file format is supported"""
        return file_path.suffix.lower() in FileParser.PARSERS
    
    @staticmethod
    def get_supported_formats() -> list:
        """Get list of supported file formats"""
        return list(FileParser.PARSERS.keys())
